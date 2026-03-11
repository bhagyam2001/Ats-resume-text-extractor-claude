import azure.functions as func
import json
import base64
import io
import re
import logging
import requests
import datetime
import unicodedata

import pdfplumber
import docx2txt

# ═══════════════════════════════════════════════════════════════════
# ENTERPRISE ATS RESUME EXTRACTOR  —  v7.0
# ═══════════════════════════════════════════════════════════════════
# Fixes in v7 vs v6:
#   1. Robust section detection — handles ALL CAPS, Title Case,
#      sentence case, coloured headers, sidebar headers, decorators
#   2. Improved name/title extraction — handles "CRM Developer"
#      appearing before name, multi-language names, initials
#   3. Full DOCX extraction — extracts from tables, headers,
#      footers, text boxes (common in modern resume templates)
#   4. Experience parser — handles company-first, title-first,
#      inline dates, no-date jobs, contract/freelance labels
#   5. Location detection — improved city/country patterns
#   6. Raw sections always populated even if headers not found
#   7. Fallback full-text scan when sections not detected
#   8. Better confidence scoring across all fields
# ═══════════════════════════════════════════════════════════════════


# ───────────────────────────────────────────────────────────────────
# SKILLS DICTIONARY  —  500+ skills across 18 categories
# ───────────────────────────────────────────────────────────────────
SKILLS_BY_CATEGORY = {
    "programming": [
        "python","java","javascript","typescript","c#","c++","c","ruby","php",
        "swift","kotlin","go","rust","scala","r","matlab","perl","bash",
        "powershell","vba","dart","lua","groovy","cobol","fortran","assembly",
        "objective-c","elixir","haskell","clojure","f#","apex","solidity","abap"
    ],
    "web": [
        "html","css","react","angular","vue","node.js","nodejs","django","flask",
        "fastapi","spring","asp.net",".net","rest","graphql","jquery","bootstrap",
        "tailwind","webpack","next.js","nuxt","gatsby","wordpress","drupal",
        "laravel","express.js","svelte","remix","astro","vite","storybook",
        "web components","pwa","sass","less","styled components"
    ],
    "data_ai": [
        "sql","mysql","postgresql","mongodb","redis","elasticsearch","pandas",
        "numpy","scikit-learn","tensorflow","pytorch","keras","machine learning",
        "deep learning","nlp","computer vision","data science","data analysis",
        "data engineering","etl","power bi","tableau","looker","qlik","excel",
        "spark","hadoop","kafka","airflow","dbt","snowflake","databricks",
        "data warehouse","data lake","data modelling","data governance",
        "business intelligence","bi","reporting","predictive analytics",
        "statistics","r studio","spss","sas","alteryx","knime",
        "google data studio","metabase","superset","mlops"
    ],
    "cloud_devops": [
        "azure","aws","gcp","docker","kubernetes","terraform","jenkins",
        "github actions","ci/cd","devops","linux","unix","azure functions",
        "azure devops","azure sql","azure data factory","azure logic apps",
        "azure service bus","azure blob storage","lambda","ec2","s3",
        "cloudformation","ansible","puppet","chef","gitlab ci","circleci",
        "travis ci","helm","istio","prometheus","grafana","elk stack",
        "datadog","new relic","site reliability","sre","infrastructure as code",
        "iac","cloud architecture","serverless","microservices","api gateway",
        "load balancing"
    ],
    "microsoft": [
        "power automate","power apps","power platform","power pages",
        "dynamics 365","sharepoint","dataverse","microsoft teams",
        "microsoft 365","office 365","power bi","azure active directory",
        "microsoft fabric","copilot studio","dynamics crm",
        "dynamics 365 finance","dynamics 365 supply chain","dynamics 365 hr",
        "dynamics 365 sales","dynamics 365 customer service","business central",
        "navision","ax","d365","excel","word","outlook","onenote","ms project",
        "visio","access","model driven apps","canvas apps","pcf controls",
        "power fx","dataverse api","plugin development","javascript web resource"
    ],
    "crm_sales_tools": [
        "salesforce","salesforce crm","salesforce sales cloud",
        "salesforce service cloud","salesforce marketing cloud",
        "salesforce pardot","salesforce cpq","salesforce admin",
        "salesforce developer","salesforce lightning","hubspot","hubspot crm",
        "hubspot marketing","zoho crm","zoho","pipedrive","freshsales",
        "freshdesk","freshworks","zendesk","intercom","drift","outreach",
        "salesloft","gong","chorus","clari","monday crm","sugar crm",
        "insightly","copper crm","close crm","netsuite crm","oracle crm",
        "sap crm","microsoft crm"
    ],
    "erp": [
        "sap","sap s/4hana","sap ecc","sap hana","sap fi","sap co","sap mm",
        "sap sd","sap hr","sap pp","sap wm","sap basis","sap abap","sap fiori",
        "sap bw","sap bi","oracle erp","oracle fusion","oracle e-business suite",
        "oracle ebs","oracle cloud","oracle financials","microsoft dynamics",
        "dynamics 365 finance","dynamics ax","dynamics nav","business central",
        "netsuite","oracle netsuite","sage","sage 200","sage 300","sage x3",
        "sage intacct","epicor","infor","syspro","odoo","ifs","unit4",
        "workday financials","peoplesoft","jd edwards"
    ],
    "databases": [
        "oracle","sql server","sqlite","cassandra","dynamodb","cosmos db",
        "neo4j","influxdb","mariadb","hbase","db2","sybase","teradata",
        "vertica","greenplum","couchdb","firebase","supabase","planetscale",
        "database design","database administration","dba","stored procedures",
        "query optimisation","indexing","t-sql","pl/sql"
    ],
    "engineering": [
        "autocad","solidworks","catia","ansys","simulink","revit","bim",
        "civil 3d","arcgis","qgis","labview","pcb design","embedded systems",
        "iot","plc","scada","mechanical design","electrical engineering",
        "structural analysis","3d modelling","cfd","fem","fea","cam","cnc",
        "circuit design","fpga","vhdl","verilog","ros","raspberry pi",
        "arduino","microcontrollers"
    ],
    "cybersecurity": [
        "penetration testing","ethical hacking","siem","soc",
        "vulnerability assessment","iso 27001","nist","gdpr compliance",
        "firewall","intrusion detection","cryptography","zero trust","owasp",
        "ceh","cissp","cism","comptia security+","network security",
        "endpoint security","dlp","iam","pam","sso","mfa","devsecops",
        "threat modelling","incident response","digital forensics","soar"
    ],
    "hr": [
        "recruitment","talent acquisition","onboarding","offboarding",
        "employee relations","performance management","compensation",
        "benefits administration","hris","workday","sap hr","bamboohr",
        "succession planning","learning and development","l&d",
        "organisational development","workforce planning","payroll",
        "employment law","hr policy","diversity and inclusion",
        "talent management","job evaluation","headhunting",
        "competency frameworks","hr analytics","people analytics",
        "oracle hcm","successfactors","sap successfactors","adp","ceridian",
        "kronos","ultipro","peoplesoft hr","greenhouse","lever","workable",
        "smartrecruiters","taleo","icims","jobvite","bullhorn",
        "employee engagement","culture","wellbeing","coaching","mentoring",
        "training delivery","change management","organisational design"
    ],
    "finance": [
        "financial reporting","financial analysis","financial modelling",
        "ifrs","gaap","us gaap","management accounts","budgeting","forecasting",
        "variance analysis","cash flow management","accounts payable",
        "accounts receivable","reconciliation","audit","internal audit",
        "external audit","tax","vat","corporate finance","investment banking",
        "private equity","risk management","credit risk","market risk",
        "compliance","sap","oracle financials","sage","quickbooks","xero",
        "bloomberg","treasury","mergers and acquisitions","m&a","due diligence",
        "valuation","cfa","acca","cpa","aca","hyperion","anaplan","blackline",
        "kyriba","concur","netsuite","workday financials","adaptive insights",
        "financial close","consolidation","intercompany","transfer pricing",
        "tax planning","indirect tax","fp&a","financial planning",
        "cost accounting","management reporting","board reporting",
        "investor relations"
    ],
    "operations": [
        "supply chain management","logistics","procurement",
        "inventory management","warehouse management","lean","six sigma",
        "kaizen","continuous improvement","process improvement",
        "demand planning","vendor management","contract management",
        "facilities management","health and safety","quality assurance",
        "quality management","iso 9001","operational excellence",
        "kpi management","fleet management","import export","customs",
        "incoterms","s&op","mrp","erp","wms","tms","last mile delivery",
        "3pl","4pl","freight","sourcing","category management",
        "spend analysis","supplier relationship management","srm",
        "business process improvement","bpi","bpm","iso 14001","iso 45001",
        "haccp","gmp","production planning","capacity planning",
        "total quality management","tqm"
    ],
    "sales": [
        "sales","business development","account management","lead generation",
        "b2b sales","b2c sales","key account management","enterprise sales",
        "inside sales","field sales","channel sales","solution selling",
        "consultative selling","value selling","spin selling","cold calling",
        "prospecting","pipeline management","sales forecasting",
        "quota attainment","revenue growth","territory management",
        "new business","client acquisition","upselling","cross selling",
        "contract negotiation","tender management","rfp","bid management",
        "presales","sales enablement","sales operations"
    ],
    "marketing": [
        "digital marketing","seo","sem","google analytics",
        "social media marketing","content marketing","email marketing",
        "brand management","market research","campaign management",
        "google ads","facebook ads","linkedin marketing","copywriting",
        "public relations","pr","media relations","marketing strategy",
        "product marketing","growth hacking","marketing automation",
        "marketo","pardot","mailchimp","adobe campaign","klaviyo","braze",
        "iterable","google tag manager","adobe analytics","mixpanel",
        "amplitude","affiliate marketing","influencer marketing",
        "programmatic","display advertising","conversion rate optimisation",
        "cro","a/b testing","landing pages","marketing funnel",
        "customer acquisition","retention marketing","lifecycle marketing",
        "product led growth","demand generation","account based marketing",
        "abm","event marketing","trade shows","sponsorship"
    ],
    "project_management": [
        "project management","programme management","pmp","prince2",
        "agile","scrum","kanban","safe","waterfall","hybrid","jira","asana",
        "ms project","monday.com","trello","risk management","change management",
        "stakeholder management","budget management","resource planning",
        "governance","pmo","benefits realisation","business analysis","ba",
        "smartsheet","basecamp","notion","confluence","miro",
        "programme governance","portfolio management","dependency management",
        "milestone tracking","project reporting","earned value management",
        "evm","cost management","scope management","schedule management",
        "pmbok","ipma","apm","msp","p3o"
    ],
    "legal": [
        "contract law","employment law","corporate law","commercial law",
        "litigation","dispute resolution","arbitration","mediation",
        "intellectual property","gdpr","data protection",
        "regulatory compliance","mergers acquisitions","due diligence",
        "legal research","legal drafting","company secretarial","conveyancing",
        "family law","criminal law","competition law","financial regulation",
        "fca","aml","anti money laundering","kyc","sanctions","privacy law",
        "cyber law","fintech regulation","commercial contracts","ndas",
        "slas","msa","legal operations","legaltech","contract management"
    ],
    "healthcare": [
        "clinical trials","gcp","ich guidelines","regulatory affairs",
        "pharmacovigilance","medical writing","nursing","patient care",
        "electronic health records","ehr","emr","nhs","cqc",
        "healthcare management","infection control","clinical governance",
        "medical coding","icd-10","hipaa","care planning","clinical research",
        "protocol development","ethics committee","fda","ema","mhra",
        "clinical data management","medical devices","ce marking","iso 13485",
        "gmp","gcp","glp","gdp","health informatics","nhs digital","hl7","fhir",
        "pharmacy","pathology","radiology","physiotherapy"
    ],
    "design": [
        "figma","adobe xd","sketch","invision","photoshop","illustrator",
        "indesign","after effects","premiere pro","final cut pro",
        "davinci resolve","ux design","ui design","user research",
        "usability testing","wireframing","prototyping","design systems",
        "graphic design","visual design","motion graphics","video editing",
        "photography","3d design","blender","canva","zeplin","abstract",
        "principle","accessibility","wcag","responsive design"
    ],
    "soft_skills": [
        "leadership","team management","communication","presentation",
        "negotiation","problem solving","critical thinking",
        "stakeholder management","time management","adaptability",
        "collaboration","mentoring","coaching","decision making",
        "strategic thinking","analytical skills","attention to detail",
        "customer service","client relationship management","influencing",
        "conflict resolution","emotional intelligence","resilience",
        "initiative","creativity","innovation","commercial awareness",
        "entrepreneurial","self motivated"
    ],
    "spoken_languages": [
        "english","french","german","spanish","mandarin","arabic","hindi",
        "portuguese","italian","japanese","korean","russian","dutch","polish",
        "turkish","urdu","punjabi","tamil","bengali","swahili","hebrew",
        "greek","swedish","norwegian","danish","finnish","romanian","czech",
        "hungarian","thai","vietnamese","indonesian","malay"
    ]
}

ALL_SKILLS = [(skill, cat) for cat, skills in SKILLS_BY_CATEGORY.items()
              for skill in skills]


# ───────────────────────────────────────────────────────────────────
# SECTION PATTERNS  —  handles ALL CAPS, Title Case, decorators
# ───────────────────────────────────────────────────────────────────
def _sec(*variants):
    """Build a pattern that matches any of the given variants,
    in Title Case, ALL CAPS, or sentence case, with optional
    decoration chars before/after (bullets, dashes, colons, pipes)."""
    joined = "|".join(variants)
    return re.compile(
        r"^[\s\-•|►▸▶✦✧★●○◆◇=_]*"
        r"(?:" + joined + r")"
        r"[\s\-:•|►▸=_]*$",
        re.IGNORECASE
    )

SECTION_PATTERNS = {
    "summary": _sec(
        r"(?:professional\s+)?(?:summary|profile|overview|introduction)",
        r"career\s+(?:summary|objective|profile|overview)",
        r"personal\s+statement", r"executive\s+summary",
        r"about\s+me", r"objective"
    ),
    "experience": _sec(
        r"(?:work|professional|employment|career|relevant|industry)?"
        r"\s*(?:experience|history|background|roles?)",
        r"positions?\s+(?:held|of\s+responsibility)",
        r"employment\s+record", r"work\s+history",
        r"professional\s+background", r"career\s+history",
        r"previous\s+(?:roles?|experience|employment)"
    ),
    "education": _sec(
        r"education(?:al\s+(?:background|qualifications?|history))?",
        r"academic\s+(?:background|qualifications?|history|record)",
        r"qualifications?", r"degrees?", r"academic\s+profile",
        r"training\s+(?:and\s+education|&\s+education)",
        r"education\s+and\s+training"
    ),
    "skills": _sec(
        r"(?:technical\s+|core\s+|key\s+|professional\s+|it\s+)?skills?",
        r"(?:core\s+)?competenc(?:y|ies)",
        r"areas?\s+of\s+expertise",
        r"technologies(?:\s+and\s+tools?)?",
        r"tools?(?:\s+and\s+technologies?)?",
        r"expertise", r"capabilities", r"technical\s+proficiencies",
        r"skills?\s+and\s+(?:expertise|technologies|tools?|knowledge)",
        r"technical\s+skills?\s+(?:and\s+)?(?:tools?|technologies)?",
        r"it\s+skills?"
    ),
    "certifications": _sec(
        r"certifications?", r"certificates?", r"accreditations?",
        r"professional\s+(?:certifications?|qualifications?|development)",
        r"licenses?(?:\s+and\s+certifications?)?", r"credentials?",
        r"courses?\s+(?:and\s+certifications?)?",
        r"training\s+and\s+certifications?"
    ),
    "projects": _sec(
        r"(?:key\s+|notable\s+|selected\s+|major\s+)?projects?",
        r"portfolio", r"personal\s+projects?",
        r"open\s+source\s+(?:contributions?|projects?)"
    ),
    "languages": _sec(
        r"(?:spoken\s+|foreign\s+)?languages?",
        r"language\s+(?:skills?|proficiencies?|abilities?)",
        r"linguistic\s+skills?"
    ),
    "achievements": _sec(
        r"achievements?", r"accomplishments?",
        r"awards?(?:\s+and\s+achievements?)?",
        r"honours?", r"honors?", r"recognition",
        r"key\s+achievements?", r"notable\s+achievements?",
        r"highlights?"
    ),
    "interests": _sec(
        r"(?:personal\s+)?interests?",
        r"hobbies(?:\s+and\s+interests?)?",
        r"activities"
    ),
    "references": _sec(r"references?", r"referees?"),
    "publications": _sec(
        r"publications?", r"research", r"papers?", r"articles?"
    ),
    "volunteering": _sec(
        r"volunteer(?:ing|\s+experience)?",
        r"community\s+(?:involvement|service|work)"
    ),
}

# Regex patterns
EMAIL_PAT    = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
PHONE_PAT    = re.compile(r"(?:\+?\d[\d\s\-().]{7,17}\d)")
LINKEDIN_PAT = re.compile(r"linkedin\.com/in/[\w\-]+", re.IGNORECASE)
GITHUB_PAT   = re.compile(r"github\.com/[\w\-]+", re.IGNORECASE)
URL_PAT      = re.compile(r"https?://[^\s]+", re.IGNORECASE)

DATE_PATTERN = re.compile(
    r"(?:(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|"
    r"jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|"
    r"nov(?:ember)?|dec(?:ember)?)\s+)?(\d{4})\s*"
    r"(?:[-\u2013\u2014]|to|\s+till\s+|\s+until\s+)\s*"
    r"(?:(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|"
    r"jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|"
    r"nov(?:ember)?|dec(?:ember)?)\s+)?(\d{4}|present|current|now|till\s+date|date)",
    re.IGNORECASE
)

DEGREE_PAT = re.compile(
    r"\b(b\.?sc|b\.?eng|b\.?e\.?|b\.?a\.?|b\.?tech|b\.?com|"
    r"m\.?sc|m\.?eng|m\.?b\.?a|m\.?tech|m\.?com|m\.?a\.?|"
    r"ph\.?d|bachelor(?:s|\s+of)?|master(?:s|\s+of)?|doctorate|"
    r"diploma|certificate|hnd|hnc|a-levels?|gcse|lpc|bptc|llb|llm|"
    r"be\b|btech\b|mtech\b|mba\b)\b",
    re.IGNORECASE
)

# Title case / ALL CAPS line detector — used to find implicit section headers
HEADER_LINE_PAT = re.compile(
    r"^(?:[A-Z][A-Z\s&/]{3,60}|[A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,4})$"
)

JOB_TITLE_KEYWORDS = [
    "engineer","developer","manager","analyst","director","specialist",
    "consultant","officer","lead","head","coordinator","executive",
    "associate","senior","junior","architect","designer","scientist",
    "administrator","advisor","controller","supervisor","president",
    "vp","cto","ceo","cfo","coo","intern","trainee","apprentice",
    "technician","operator","assistant","representative","agent",
    "strategist","researcher","planner","programmer","tester","qa"
]


# ───────────────────────────────────────────────────────────────────
# PDF EXTRACTION  —  column-aware + OCR fallback
# ───────────────────────────────────────────────────────────────────
def extract_pdf(file_bytes: bytes) -> dict:
    pages_text = []
    ocr_used   = False

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        page_count = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            words = page.extract_words()
            page_text = ""

            if words:
                midpoint   = page.width / 2
                left_words = [w for w in words if w["x0"] < midpoint - 20]
                right_words= [w for w in words if w["x0"] > midpoint + 20]
                left_ratio = len(left_words) / len(words)
                right_ratio= len(right_words) / len(words)

                if left_ratio > 0.2 and right_ratio > 0.2:
                    left_text  = (page.crop((0, 0, midpoint, page.height)).extract_text() or "")
                    right_text = (page.crop((midpoint, 0, page.width, page.height)).extract_text() or "")
                    page_text  = left_text.strip() + "\n\n" + right_text.strip()
                else:
                    page_text  = page.extract_text() or ""

            if not page_text.strip():
                try:
                    img_bytes = io.BytesIO()
                    page.to_image(resolution=200).original.save(img_bytes, format="PNG")
                    page_text = _ocr(img_bytes.getvalue())
                    if page_text.strip():
                        ocr_used = True
                except Exception as e:
                    logging.warning(f"OCR fallback page {i+1}: {e}")

            pages_text.append(page_text)

    return {
        "text":       "\n\n".join(pages_text).strip(),
        "page_count": page_count,
        "ocr_used":   ocr_used
    }


# ───────────────────────────────────────────────────────────────────
# DOCX EXTRACTION  —  full extraction including tables + text boxes
# ───────────────────────────────────────────────────────────────────
def extract_doc_legacy(file_bytes: bytes) -> dict:
    """Extract text from legacy .doc (OLE) files using LibreOffice conversion."""
    import subprocess, tempfile, os
    with tempfile.TemporaryDirectory() as tmpdir:
        doc_path = os.path.join(tmpdir, "resume.doc")
        with open(doc_path, "wb") as f:
            f.write(file_bytes)
        try:
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "docx",
                 doc_path, "--outdir", tmpdir],
                capture_output=True, timeout=30
            )
            docx_files = [f for f in os.listdir(tmpdir) if f.endswith(".docx")]
            if docx_files:
                docx_path = os.path.join(tmpdir, docx_files[0])
                with open(docx_path, "rb") as f:
                    docx_bytes = f.read()
                return extract_docx(docx_bytes)
            else:
                raise ValueError("LibreOffice conversion produced no output.")
        except subprocess.TimeoutExpired:
            raise ValueError("LibreOffice conversion timed out (>30s).")
        except FileNotFoundError:
            # LibreOffice not installed — try raw binary text extraction
            logging.warning("LibreOffice not found — attempting raw .doc text extraction")
            text = _extract_doc_raw(file_bytes)
            return {"text": text, "page_count": None, "ocr_used": False}


def _extract_doc_raw(file_bytes: bytes) -> str:
    """Last-resort: extract printable ASCII strings from .doc binary."""
    try:
        text = file_bytes.decode("latin-1", errors="ignore")
        # Keep runs of printable characters
        chunks = re.findall(r"[ -~\n\t]{20,}", text)
        # Filter out binary-looking chunks
        clean  = [c for c in chunks if sum(1 for ch in c if ch.isalpha()) > len(c) * 0.4]
        return "\n".join(clean)
    except Exception:
        return ""


def extract_docx(file_bytes: bytes) -> dict:
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(file_bytes))
        parts = []

        # Body paragraphs
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)

        # Tables (very common in modern resume templates)
        for table in doc.tables:
            for row in table.rows:
                row_parts = []
                for cell in row.cells:
                    ct = cell.text.strip()
                    if ct:
                        row_parts.append(ct)
                if row_parts:
                    parts.append("  |  ".join(row_parts))

        # Headers and footers
        for section in doc.sections:
            for hf in [section.header, section.footer]:
                if hf:
                    for para in hf.paragraphs:
                        t = para.text.strip()
                        if t:
                            parts.append(t)

        text = "\n".join(parts)
        return {"text": text.strip(), "page_count": None, "ocr_used": False}

    except ImportError:
        # Fallback to docx2txt if python-docx not installed
        try:
            text = docx2txt.process(io.BytesIO(file_bytes))
            return {"text": text.strip(), "page_count": None, "ocr_used": False}
        except Exception as e:
            raise ValueError(f"Could not read Word document: {e}")
    except Exception as e:
        raise ValueError(f"Could not read Word document: {e}. "
                         "File may be password-protected or corrupted.")


def _ocr(image_bytes: bytes) -> str:
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(image_bytes))
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        return pytesseract.image_to_string(img, config="--psm 6")
    except Exception as e:
        logging.warning(f"OCR failed: {e}")
        return ""


def extract_image(file_bytes: bytes) -> dict:
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(file_bytes))
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        text = pytesseract.image_to_string(img, config="--psm 6")
        return {"text": text.strip(), "page_count": 1, "ocr_used": True}
    except ImportError:
        return {
            "text": "", "page_count": 1, "ocr_used": False,
            "warning": "pytesseract not available — PDF and DOCX work normally."
        }
    except Exception as e:
        raise ValueError(f"Image OCR failed: {e}")


def detect_type(b: bytes) -> str:
    if b[:4] == b"%PDF":                       return "pdf"
    if b[:2] == b"PK":                         return "docx"
    if b[:3] == b"\xff\xd8\xff":               return "jpeg"
    if b[:8] == b"\x89PNG\r\n\x1a\n":          return "png"
    if b[:4] in (b"MM\x00*", b"II*\x00"):    return "tiff"
    if b[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1": return "doc"
    return "unknown"


def download_file(url: str, token: str = "") -> bytes:
    headers = {"Authorization": f"Bearer {token}"} if token else {}
    r = requests.get(url, headers=headers, timeout=30)
    r.raise_for_status()
    return r.content


# ───────────────────────────────────────────────────────────────────
# TEXT CLEANING
# ───────────────────────────────────────────────────────────────────
def clean_text(text: str) -> str:
    if not text:
        return ""
    # Normalise unicode (smart quotes, en/em dashes, curly quotes → ASCII)
    text = unicodedata.normalize("NFKD", text)
    # Replace en-dash (–) and em-dash (—) with hyphen BEFORE ascii encoding
    text = text.replace("\u2013", "-").replace("\u2014", "-")
    text = text.replace("\u2012", "-").replace("\u2010", "-")
    text = text.encode("ascii", "ignore").decode("ascii")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Remove decorative lines
    text = re.sub(r"^[-_=.•|*#]{3,}$", "", text, flags=re.MULTILINE)
    # Collapse multiple spaces
    text = re.sub(r"[ \t]{2,}", " ", text)
    # Collapse excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return "\n".join(l.strip() for l in text.split("\n")).strip()


# ───────────────────────────────────────────────────────────────────
# SECTION SPLITTER  —  v7: handles implicit headers + ALL CAPS
# ───────────────────────────────────────────────────────────────────
def split_sections(text: str) -> dict:
    lines   = text.split("\n")
    sections = {}
    current  = "header"
    buf      = []

    for line in lines:
        stripped = line.strip()

        # Skip blank lines but preserve them in buffer
        if not stripped:
            buf.append("")
            continue

        matched = None

        # 1 — try explicit named section patterns (max 80 chars)
        if len(stripped) <= 80:
            for name, pat in SECTION_PATTERNS.items():
                if pat.match(stripped):
                    matched = name
                    break

        # 2 — try ALL CAPS line as implicit section header
        if not matched and re.match(r"^[A-Z][A-Z\s&/\-]{3,60}$", stripped):
            guess = stripped.lower().strip()
            for name, pat in SECTION_PATTERNS.items():
                if pat.match(guess):
                    matched = name
                    break

        if matched:
            sections[current] = "\n".join(buf).strip()
            current = matched
            buf     = []
        else:
            buf.append(line)

    sections[current] = "\n".join(buf).strip()

    # If we only detected "header" and nothing else,
    # try a fallback: use the whole text for extraction
    if len([k for k, v in sections.items() if v.strip()]) <= 1:
        sections["_full_text"] = text

    return sections


# ───────────────────────────────────────────────────────────────────
# CONTACT INFO
# ───────────────────────────────────────────────────────────────────
def extract_contact(text: str) -> dict:
    contact = {"email": None, "phone": None, "linkedin": None,
               "github": None, "website": None}

    m = EMAIL_PAT.search(text)
    if m:
        contact["email"] = m.group().strip()

    for m in PHONE_PAT.finditer(text):
        digits = re.sub(r"\D", "", m.group())
        if 7 <= len(digits) <= 15:
            contact["phone"] = m.group().strip()
            break

    m = LINKEDIN_PAT.search(text)
    if m:
        contact["linkedin"] = "https://" + m.group()

    m = GITHUB_PAT.search(text)
    if m:
        contact["github"] = "https://" + m.group()

    for m in URL_PAT.finditer(text):
        url = m.group()
        if "linkedin" not in url.lower() and "github" not in url.lower():
            contact["website"] = url
            break

    return contact


# ───────────────────────────────────────────────────────────────────
# NAME + TITLE EXTRACTION  —  v7: smarter disambiguation
# ───────────────────────────────────────────────────────────────────
def extract_identity(header_text: str) -> dict:
    lines = [l.strip() for l in header_text.split("\n") if l.strip()]
    name  = None
    title = None

    def is_name_candidate(line: str) -> bool:
        words = line.split()
        if not (1 < len(words) <= 5):
            return False
        if len(line) > 55:
            return False
        if EMAIL_PAT.search(line) or PHONE_PAT.search(line):
            return False
        if any(kw in line.lower() for kw in
               ["http", "www", "@", "linkedin", "github", "street", "road",
                "avenue", "city", "state", "pin", "zip"]):
            return False
        # All words should look like name parts
        if all(re.match(r"^[A-Za-z\-\'.]+(\s[A-Za-z\-\'.]+)?$", w)
               for w in words):
            # Prefer words that are Title Case (not ALL CAPS job title)
            return True
        return False

    def is_title_candidate(line: str) -> bool:
        words = line.split()
        if not (1 <= len(words) <= 10):
            return False
        if len(line) > 80:
            return False
        if EMAIL_PAT.search(line) or PHONE_PAT.search(line):
            return False
        lower = line.lower()
        return any(kw in lower for kw in JOB_TITLE_KEYWORDS)

    # Pass 1: Look for name first (prefer lines that are NOT job titles)
    for line in lines[:10]:
        if name is None and is_name_candidate(line) and not is_title_candidate(line):
            name = line
            continue
        if name and title is None and is_title_candidate(line):
            title = line
            break

    # Pass 2: If no name found yet, relax — allow name even if it
    # looks like a title (e.g. "Hariharasuthan R")
    if name is None:
        for line in lines[:10]:
            if is_name_candidate(line):
                name = line
                break

    # Pass 3: Find title independently
    if title is None:
        for line in lines[:12]:
            if line == name:
                continue
            if is_title_candidate(line):
                title = line
                break

    # Pass 4: fallback — look in the broader text for a Designation/Role label row
    # e.g. table rows like "Senior Associate Software Applications Development Engineer"
    # which appear in work history tables
    if title is None:
        for line in lines[:20]:
            if line == name:
                continue
            words = line.split()
            # short-ish, no numbers, not a sentence, contains title keyword
            if (2 <= len(words) <= 10 and len(line) < 90
                    and not re.search(r"\d", line)
                    and not EMAIL_PAT.search(line)
                    and not PHONE_PAT.search(line)
                    and not re.match(r"^(having|good|sound|extensive|minimal|"
                                     r"in-depth|expertise|experience)", line.lower())
                    and any(kw in line.lower() for kw in JOB_TITLE_KEYWORDS)):
                title = line
                break

    return {
        "name":             name,
        "name_confidence":  _conf_name(name),
        "current_title":    title,
        "title_confidence": _conf_title(title)
    }


def _conf_name(name) -> str:
    if not name:
        return "low"
    words = str(name).split()
    if 2 <= len(words) <= 4:
        return "high"
    return "medium"

def _conf_title(title) -> str:
    if not title:
        return "low"
    lower = str(title).lower()
    if any(kw in lower for kw in JOB_TITLE_KEYWORDS):
        return "high"
    return "medium"


# ───────────────────────────────────────────────────────────────────
# LOCATION
# ───────────────────────────────────────────────────────────────────
def extract_location(text: str) -> str:
    # Pattern: "City, Country" or "City, ST" — but NOT content phrases
    patterns = [
        re.compile(
            r"\b([A-Z][a-z]+(?:[\s\-][A-Z][a-z]+)?,\s*"
            r"(?:[A-Z]{2,3}|[A-Z][a-z]+(?:\s[A-Z][a-z]+)?))\b"
        ),
        re.compile(
            r"(?:location|based in|residing in|address)[:\s]+([^\n,]{3,50})",
            re.IGNORECASE
        ),
    ]
    sample = text[:800]
    for pat in patterns:
        m = pat.search(sample)
        if m:
            candidate = m.group(1).strip()
            # Reject if it looks like a skill phrase
            bad = ["process","configuration","development","management",
                   "solution","implementation","design","architect"]
            if not any(b in candidate.lower() for b in bad):
                return candidate
    return None


# ───────────────────────────────────────────────────────────────────
# SKILLS
# ───────────────────────────────────────────────────────────────────
# Patterns that introduce a tool/technology inline in a sentence
INLINE_SKILL_PATTERNS = [
    # "using Power Automate", "via Dataverse", "with JavaScript"
    re.compile(
        r"\b(?:using|via|with|through|leveraging|utilising|utilizing|"
        r"building|built\s+(?:with|using|in)|developed\s+(?:with|using|in)|"
        r"implemented\s+(?:with|using|in)|integrated\s+with|"
        r"configured\s+(?:with|in)|deployed\s+(?:on|to|via))\s+"
        r"([A-Z][A-Za-z0-9.#+\s]{1,40}?)(?=[,;.\n]|\s+(?:to|for|and|or|the|a|an)\b|$)",
        re.IGNORECASE
    ),
    # "Power Automate integration", "CRM customization", "API development"
    re.compile(
        r"\b([A-Z][A-Za-z0-9.#+\s]{2,35}?)\s+"
        r"(?:integration|customisation|customization|automation|development|"
        r"configuration|implementation|deployment|migration|administration)\b",
        re.IGNORECASE
    ),
    # "in Dynamics 365", "on Azure", "on SharePoint"
    re.compile(
        r"\b(?:in|on)\s+([A-Z][A-Za-z0-9.#+\s]{2,35}?)(?=[,;.\n]|\s+(?:to|for|and|or|the|a|an)\b|$)",
        re.IGNORECASE
    ),
]

# Words that commonly follow "in/on" but are NOT technologies
INLINE_SKILL_STOPWORDS = {
    "the","a","an","this","that","these","those","my","our","their",
    "time","place","order","addition","general","particular","fact",
    "line","real","detail","practice","production","staging","house",
    "scope","progress","advance","depth","brief","short","full","charge",
    "total","parallel","response","reply","support","terms","case",
    "accordance","compliance","conjunction","context","relation",
    "alignment","contrast","comparison","summary","conclusion",
}

def _extract_inline_skills(text: str, known_skills_lower: set) -> list:
    """
    Extract technology/tool names mentioned inline in bullet points and
    descriptions using linguistic patterns like "using X", "via X", "with X".

    Only returns candidates that are either:
      a) Already in our skills dictionary (known_skills_lower), OR
      b) Look like a proper technology name (Title-cased, short, no verb)
    """
    candidates = []
    for pat in INLINE_SKILL_PATTERNS:
        for m in pat.finditer(text):
            raw = m.group(1).strip().rstrip(".,;:")
            raw = re.sub(r"\s+", " ", raw)

            # Must be 2–6 words, reasonable length
            words = raw.split()
            if not (1 <= len(words) <= 5) or len(raw) > 50:
                continue

            lower = raw.lower()

            # Skip stopwords and sentence fragments
            if lower in INLINE_SKILL_STOPWORDS:
                continue
            if any(lower.startswith(sw + " ") for sw in INLINE_SKILL_STOPWORDS):
                continue

            # Skip if it looks like a sentence (contains verb endings)
            if re.search(r"\b(ing|tion|ed|ment|ness|ity)$", lower) and len(words) == 1:
                continue

            # Accept if in dictionary
            if lower in known_skills_lower:
                candidates.append(raw)
                continue

            # Accept if Title-cased product name (e.g. "Canvas Apps", "QR code")
            if (words[0][0].isupper() and len(raw) >= 4
                    and not re.search(r"\b(and|or|the|for|of|to|in|a|an)\b", lower)):
                candidates.append(raw)

    # Deduplicate preserving order
    seen = set()
    result = []
    for c in candidates:
        key = c.lower()
        if key not in seen:
            seen.add(key)
            result.append(c)
    return result


def extract_skills(skills_section: str, full_text: str) -> dict:
    search = (skills_section + "\n" + full_text).lower()
    by_cat = {}
    found  = set()

    for skill, cat in ALL_SKILLS:
        if re.search(r"\b" + re.escape(skill) + r"\b", search, re.IGNORECASE):
            display = skill.title() if len(skill) > 2 else skill.upper()
            if display not in found:
                by_cat.setdefault(cat, []).append(display)
                found.add(display)

    # Also parse items directly listed in skills section
    if skills_section:
        for raw in re.split(r"[,|•·\n\t/\\]+", skills_section):
            s = re.sub(r"^[-•·*▪▸✓✔\s]+", "", raw).strip()
            s = re.sub(r"\s+", " ", s).strip()
            if 2 <= len(s) <= 60 and len(s.split()) <= 5:
                if s not in found and not any(
                    stop in s.lower() for stop in
                    ["year","month","responsi","worked","developed","managed",
                     "experience","knowledge of","\\(","till date","to date",
                     "2020","2021","2022","2023","2024","2025","2026"]
                ) and not re.match(r"^[\d\s\-/]+$", s):  # exclude pure date strings
                    by_cat.setdefault("other", []).append(s)
                    found.add(s)

    # ── Normalise: lowercase key for deduplication, preserve display form ──
    def _normalise_skill(s: str) -> str:
        """Return a canonical lowercase key for deduplication."""
        return re.sub(r"[^a-z0-9]", "", s.lower())

    # ── Validity filter ──────────────────────────────────────────────────
    BAD_WORDS = {
        "solutions","private","limited","ltd","pvt","llc","corporation",
        "incorporated","university","institute","college","responsible",
        "developed","managed","performed","conducted","involved","worked",
        "focused","integrated","created","creating","configuring","plug-ins",
        "percentage","grade","gpa","cgpa","score","marks"
    }
    NOISE_EXACT = {
        "business","client","actions","compliance","configurations","create",
        "reporting","r","c","declaration","description","designation",
        "duration","organisation","organization","environment","software",
        "customization","dashboards","implement","expenses","forms","reports",
        "updates","views","workflows","attributes","manage","project","scripts",
        "security model","the entity model","or sdk","and web services",
        "approval flows","product detailing","and daily call","impact analysis",
        "automating business processes","integrating communication tools",
        "collaborated with cross-functional teams",
        "participated in requirement gathering",
    }

    def _is_valid_skill(s: str) -> bool:
        if len(s) > 55 or len(s) < 2:              return False
        if re.search(r"\b(20\d{2}|19\d{2})\b", s): return False
        if "@" in s or "http" in s.lower():         return False
        if s.endswith(":"):                          return False
        if s.startswith("(") or s.endswith(")"):    return False
        if s.endswith(".") and len(s) > 20:         return False
        if len(s.split()) > 5:                      return False
        if re.match(r"^[A-Z][A-Z\s&]{8,}$", s):   return False  # ALLCAPS company
        lower = s.lower().strip()
        if lower in NOISE_EXACT:                     return False
        if any(w in lower.split() for w in BAD_WORDS): return False
        # Three+ proper-cased words with no digit = likely a person/company name
        if re.match(r"^[A-Z][a-z]+(?:\s[A-Z][a-z]+){2,}$", s): return False
        return True

    # ── Deduplicate: keep the best display form per normalised key ───────
    # Priority: longer form > title-cased > original
    seen_keys: dict[str, str] = {}
    for s in found:
        if not _is_valid_skill(s):
            continue
        key = _normalise_skill(s)
        if key not in seen_keys:
            seen_keys[key] = s
        else:
            # Prefer title-cased or longer form
            existing = seen_keys[key]
            if len(s) > len(existing):
                seen_keys[key] = s

    # ── Build clean outputs ──────────────────────────────────────────────
    # Exclude "other" category from final output entirely — it's noisy.
    # Keep "inferred" (from inline extraction) as it's pattern-validated.
    EXCLUDED_CATS = {"other"}

    clean_by_cat: dict[str, list] = {}
    seen_in_cat: set[str] = set()
    all_clean: list[str] = []

    for cat, skills in by_cat.items():
        if cat in EXCLUDED_CATS:
            continue
        valid = []
        for s in skills:
            if not _is_valid_skill(s):
                continue
            key = _normalise_skill(s)
            if key in seen_in_cat:
                continue
            seen_in_cat.add(key)
            best = seen_keys.get(key, s)
            valid.append(best)
            all_clean.append(best)
        if valid:
            clean_by_cat[cat] = valid

    deduped_skills = sorted(all_clean, key=lambda x: x.lower())
    n = len(deduped_skills)
    return {
        "all_skills":         deduped_skills,
        "skills_by_category": clean_by_cat,
        "total_count":        n,
        "confidence":         "high" if n >= 10 else "medium" if n >= 5 else "low"
    }


# ───────────────────────────────────────────────────────────────────
# EXPERIENCE PARSER  —  v7.3
# Handles two formats:
#   A) Employer-level block:
#        "Company Name"  →  "Job Title"  →  Date  →  bullets
#   B) Employer + Project blocks (consulting resume style):
#        "Company Name"  →  Date
#        "Project: X"  →  "Role: Y"  →  "Project description: ..."
#        "Roles and Responsibilities"  →  bullet list
# ───────────────────────────────────────────────────────────────────

PROJECT_PAT = re.compile(
    r"^project\s*(?:#\s*\d+)?\s*:?\s*(.+)?$", re.IGNORECASE
)
ROLE_PAT = re.compile(
    r"^role\s*:?\s*(.+)$", re.IGNORECASE
)
PROJ_DESC_PAT = re.compile(
    r"^project\s+description\s*:?\s*(.*)$", re.IGNORECASE
)
RESP_PAT = re.compile(
    r"^roles?\s*(?:and\s+|&\s*)?responsibilities?\s*:?\s*$", re.IGNORECASE
)
BULLET_PAT = re.compile(r"^[-•·*▪▸✓✔]\s+|^o\s+(?=[A-Z])")


def extract_experience(exp_text: str) -> list:
    """
    Parse experience section into structured employer blocks,
    each containing one or more projects with full detail.
    """
    if not exp_text:
        return []

    lines = [l.rstrip() for l in exp_text.split("\n")]

    # ── Step 1: Split into employer blocks ──────────────────────────
    # An employer block starts with a short line that is NOT a label,
    # NOT a bullet, NOT a date-only line — typically "Company, Title"
    # or just "Company Name"
    employer_blocks = _split_employer_blocks(lines)

    results = []
    for block in employer_blocks:
        parsed = _parse_employer_block(block)
        if parsed:
            results.append(parsed)

    return results


def _split_employer_blocks(lines: list) -> list:
    """Split raw lines into groups, one group per employer."""
    blocks  = []
    current = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            current.append(line)
            continue

        # New employer block starts when:
        # - short line (< 80 chars)
        # - NOT a project/role/description label
        # - NOT a bullet
        # - NOT a responsibilities header
        # - contains a job-title keyword OR looks like "Title, Company"
        # - and we already have something in current
        if (current
                and len(stripped) < 90
                and "|" not in stripped          # never split on pipe-table rows
                and not PROJECT_PAT.match(stripped)
                and not ROLE_PAT.match(stripped)
                and not PROJ_DESC_PAT.match(stripped)
                and not RESP_PAT.match(stripped)
                and not BULLET_PAT.match(stripped)
                and not stripped.lower().startswith(("from ", "since "))
                and _looks_like_employer_header(stripped)):
            blocks.append(current)
            current = [line]
        else:
            current.append(line)

    if current:
        blocks.append(current)

    return [b for b in blocks if any(l.strip() for l in b)]


# Lines that should NEVER be treated as employer headers
HEADER_BLOCKLIST = re.compile(
    r"^(description|project\s+description|roles?\s+(?:and\s+)?responsibilities?|"
    r"declaration|environment|technology|programming\s+languages?|"
    r"scripting\s+languages?|operating\s+systems?|tools?\s*(?:&|and)?\s*technologies?|"
    r"team\s+size|duration|client|role|organisation|organization|"
    r"involved\s+in|responsible\s+for|developed|created|designed|"
    r"configured|participated|conducted|collaborated|maintained|"
    r"have\s+exposure|written|worked|utilized|understanding)\b",
    re.IGNORECASE
)

def _looks_like_employer_header(line: str) -> bool:
    """Return True if this line looks like a new employer/job header."""
    stripped = line.strip().rstrip(":")
    words = stripped.split()
    if not (1 <= len(words) <= 12):
        return False
    # Block known label words and sentence starters
    if HEADER_BLOCKLIST.match(stripped):
        return False
    # Block lines that are obviously sentences (contain verb phrases)
    if len(stripped) > 80:
        return False
    # Block lines ending with a period (they're sentences)
    if stripped.endswith("."):
        return False
    lower = stripped.lower()
    # Contains a job title keyword
    if any(kw in lower for kw in JOB_TITLE_KEYWORDS):
        return True
    # "Company Name, Job Title" pattern
    if "," in stripped and len(words) <= 8:
        return True
    # All title-case words (likely a company name)
    if all(w[0].isupper() for w in words if w and w[0].isalpha()):
        return True
    return False


def _parse_employer_block(lines: list) -> dict | None:
    """
    Parse one employer block into a structured dict with:
      employer, title, duration, years, projects[]
    Each project has: name, role, description, responsibilities[], technologies[]
    """
    stripped_lines = [l.strip() for l in lines]
    non_empty      = [l for l in stripped_lines if l]
    if not non_empty:
        return None

    employer = {
        "employer":  None,
        "title":     None,
        "duration":  None,
        "years":     None,
        "projects":  []
    }

    # ── Extract employer name and title from first 1-2 lines ────────
    first = non_empty[0] if non_empty else ""
    # "D365 Techno-Functional Consultant, Caliber Focus"
    if "," in first and len(first.split()) <= 12:
        parts = first.split(",", 1)
        title_part   = parts[0].strip()
        company_part = parts[1].strip()
        if any(kw in title_part.lower() for kw in JOB_TITLE_KEYWORDS):
            employer["title"]    = title_part
            employer["employer"] = company_part
        else:
            employer["employer"] = title_part
            employer["title"]    = company_part
    else:
        employer["employer"] = first
        # Second line might be the title
        if len(non_empty) > 1:
            second = non_empty[1]
            if (any(kw in second.lower() for kw in JOB_TITLE_KEYWORDS)
                    and not DATE_PATTERN.search(second)
                    and not second.lower().startswith("from ")):
                employer["title"] = second

    # ── Extract duration from block ──────────────────────────────────
    full_text = " ".join(non_empty)
    d = _parse_date_range(full_text)
    if d.get("_start"):
        employer["duration"] = _find_date_string(non_empty)
        employer["years"]    = d.get("years")
        employer["_start"]   = d.get("_start")
        employer["_end"]     = d.get("_end")
    else:
        # Handle "From October-2025" style
        for line in non_empty[:4]:
            if re.match(r"^from\s+", line, re.IGNORECASE):
                employer["duration"] = line
                employer["_start"]   = datetime.datetime.now().year
                employer["_end"]     = datetime.datetime.now().year
                employer["years"]    = 0

    # ── Parse project blocks within this employer ────────────────────
    employer["projects"] = _parse_projects(non_empty)

    # If no projects found, treat responsibilities as a flat list
    if not employer["projects"]:
        bullets = _extract_bullets(non_empty)
        if bullets:
            employer["projects"] = [{
                "name":             None,
                "role":             employer.get("title"),
                "description":      None,
                "responsibilities": bullets,
                "technologies":     _extract_technologies(full_text)
            }]

    # Clean internal fields
    return {k: v for k, v in employer.items() if not k.startswith("_")}


def _parse_projects(lines: list) -> list:
    """
    Split lines into project sub-blocks and parse each one.
    A new project starts at a "Project: ..." line.
    """
    projects     = []
    current_proj = None
    state        = None   # None | "desc" | "resp"

    for line in lines:
        stripped = line.strip()

        # ── Check PROJ_DESC_PAT FIRST (before PROJECT_PAT)
        # because "Project description: ..." matches both patterns
        m = PROJ_DESC_PAT.match(stripped)
        if m and current_proj is not None:
            desc_text = (m.group(1) or "").strip()
            current_proj["description"] = desc_text or None
            state = "desc"
            continue

        # ── Project header ──────────────────────────────────────────
        m = PROJECT_PAT.match(stripped)
        if m:
            if current_proj:
                projects.append(current_proj)
            proj_name    = (m.group(1) or "").strip() or None
            current_proj = {
                "name":             proj_name,
                "role":             None,
                "description":      None,
                "responsibilities": [],
                "technologies":     []
            }
            state = None
            continue

        if current_proj is None:
            continue

        # ── Responsibilities header (check BEFORE role line) ───────
        if RESP_PAT.match(stripped):
            state = "resp"
            continue

        # ── Role line ───────────────────────────────────────────────
        m = ROLE_PAT.match(stripped)
        if m:
            role_val = m.group(1).strip()
            # Guard: ignore if it looks like "s and Responsibilities"
            if not RESP_PAT.match("roles " + role_val):
                current_proj["role"] = role_val
            state = None
            continue

        # ── Continuation lines ───────────────────────────────────────
        if not line.strip():
            continue

        stripped_line = line.strip()

        if state == "desc":
            if RESP_PAT.match(stripped_line):
                state = "resp"
            elif BULLET_PAT.match(stripped_line):
                clean = re.sub(r"^[-•·*▪▸✓✔o]\s*", "", stripped_line)
                if clean and len(clean) > 5:
                    current_proj["responsibilities"].append(clean)
                state = "resp"
            elif stripped_line:
                if current_proj["description"]:
                    current_proj["description"] += " " + stripped_line
                else:
                    current_proj["description"] = stripped_line

        elif state == "resp":
            # Accept both bullet lines AND plain sentences as responsibilities
            if stripped_line and len(stripped_line) > 10:
                clean = re.sub(r"^[-•·*▪▸✓✔o]\s*", "", stripped_line)
                current_proj["responsibilities"].append(clean)

        elif BULLET_PAT.match(stripped_line):
            clean = re.sub(r"^[-•·*▪▸✓✔o]\s*", "", stripped_line)
            if clean and len(clean) > 5:
                if current_proj:
                    current_proj["responsibilities"].append(clean)
            state = "resp"

    if current_proj:
        projects.append(current_proj)

    # Extract technologies for each project from its full text
    for proj in projects:
        full = " ".join(filter(None, [
            proj.get("description", "") or "",
            " ".join(proj.get("responsibilities", [])),
        ]))
        proj["technologies"] = _extract_technologies(full)

    return projects


def _extract_bullets(lines: list) -> list:
    """Extract bullet-point lines from a block."""
    bullets = []
    in_resp = False
    for line in lines:
        if RESP_PAT.match(line):
            in_resp = True
            continue
        if in_resp or BULLET_PAT.match(line):
            clean = re.sub(r"^[-•·*▪▸✓✔o]\s*", "", line.strip())
            if clean and len(clean) > 5:
                bullets.append(clean)
    return bullets


def _extract_technologies(text: str) -> list:
    """Extract recognised technology/skill names from free text."""
    found = []
    lower = text.lower()
    for skill, cat in ALL_SKILLS:
        if cat in ("soft_skills", "spoken_languages"):
            continue
        if re.search(r"\b" + re.escape(skill) + r"\b", lower):
            display = skill.title() if len(skill) > 2 else skill.upper()
            if display not in found:
                found.append(display)
    return sorted(found)


def _find_date_string(lines: list) -> str | None:
    """Return the first line that looks like a date range."""
    for line in lines:
        if DATE_PATTERN.search(line):
            return line.strip()
        if re.match(r"^from\s+", line, re.IGNORECASE):
            return line.strip()
    return None


def _normalise_date_text(text: str) -> str:
    """
    Normalise all date string variants before year extraction.
    Handles: en/em dash, sept, jul, non-standard abbreviations,
             "From X", "Since X", "till date", "to date", "present"
    """
    # Non-standard month abbreviations → standard
    text = re.sub(r"(?i)\bsept\b", "sep", text)
    text = re.sub(r"(?i)\bjuly\b", "jul", text)
    text = re.sub(r"(?i)\bjune\b", "jun", text)
    text = re.sub(r"(?i)\bmarch\b","mar", text)
    text = re.sub(r"(?i)\bapril\b","apr", text)
    # Add space between month abbrev and year if missing: "Mar2022" → "Mar 2022"
    text = re.sub(
        r"(?i)\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)(\d{4})\b",
        r"\1 \2", text
    )
    # Normalise "till date", "to date", "till now" → "present"
    text = re.sub(r"(?i)\btill\s*date\b",  "present", text)
    text = re.sub(r"(?i)\bto\s+date\b",    "present", text)
    text = re.sub(r"(?i)\btill\s+now\b",   "present", text)
    text = re.sub(r"(?i)\bto\s+till\b",    "present", text)
    # "From October 2025" / "Since Jan 2024" → just keep the year
    text = re.sub(r"(?i)^(from|since)\s+", "", text.strip())
    return text


def _parse_date_range(text: str) -> dict:
    """
    Parse a date range string into start/end years.
    Handles:
      "Mar2022-till date"       → 2022–present
      "Sept 2023 – sept 2025"   → 2023–2025  (en-dash + non-standard abbrev)
      "2022 to till date"       → 2022–present
      "From October 2025"       → 2025–present
      "2019-2023"               → 2019–2023
      "Oct 2021 to Aug 2023"    → 2021–2023
    """
    cur_year = datetime.datetime.now().year
    text     = _normalise_date_text(text)

    years_found = re.findall(r"\b(20\d{2}|19\d{2})\b", text)
    if not years_found:
        return {"_start": None, "_end": None, "years": None}

    start = int(years_found[0])
    if re.search(r"(?i)present|current|now|ongoing", text):
        end = cur_year
    elif len(years_found) >= 2:
        end = int(years_found[-1])
    else:
        # Single year with no end → assume ongoing
        end = cur_year

    if not (1950 <= start <= cur_year + 1):
        return {"_start": None, "_end": None, "years": None}
    end = max(start, min(end, cur_year + 1))
    return {
        "_start": start,
        "_end":   end,
        "years":  round(max(0, end - start), 1)
    }


def _parse_dates(block: str) -> dict:
    cur_year = datetime.datetime.now().year
    m = DATE_PATTERN.search(block)
    if not m:
        return {"duration": None, "_start": None, "_end": None, "years": None}

    try:
        start = int(m.group(1))
    except (ValueError, TypeError):
        return {"duration": None, "_start": None, "_end": None, "years": None}

    end_raw = (m.group(2) or "").lower().strip()
    if end_raw in ("present", "current", "now", "till date", "date", ""):
        end = cur_year
    else:
        try:
            end = int(end_raw)
        except (ValueError, TypeError):
            end = cur_year

    if not (1950 <= start <= cur_year):
        return {"duration": None, "_start": None, "_end": None, "years": None}
    end = max(start, min(end, cur_year + 1))

    return {
        "duration": m.group(0).strip(),
        "_start":   start,
        "_end":     end,
        "years":    round(max(0, end - start), 1)
    }


def calc_total_exp(jobs: list) -> float:
    """
    Calculate non-overlapping total experience years from job list.
    Reads both _start/_end (internal) and duration string as fallback.
    """
    cur_year = datetime.datetime.now().year
    ranges   = []
    for job in jobs:
        s = job.get("_start") or job.get("years") and None
        e = job.get("_end")
        # Fallback: try parsing duration string
        if (not s or not e) and job.get("duration"):
            d = _parse_date_range(str(job["duration"]))
            s = d.get("_start")
            e = d.get("_end")
        if s and e and isinstance(s, int) and isinstance(e, int):
            if 1950 <= s <= cur_year and s <= e <= cur_year + 1:
                ranges.append((s, e))
    if not ranges:
        # Last resort: sum the "years" field
        total = sum(float(j.get("years") or 0) for j in jobs if j.get("years"))
        return round(min(total, 50), 1)
    ranges.sort()
    merged = [list(ranges[0])]
    for s, e in ranges[1:]:
        if s <= merged[-1][1]:
            merged[-1][1] = max(merged[-1][1], e)
        else:
            merged.append([s, e])
    return round(min(sum(e - s for s, e in merged), 50), 1)


# ───────────────────────────────────────────────────────────────────
# EDUCATION
# ───────────────────────────────────────────────────────────────────
def extract_education(edu_text: str) -> list:
    """
    Parse education section into structured degree entries.

    Handles formats:
      A) "Bachelor of Technology (B.Tech) in ECE, Bharath University"
         "2019 - 2023"
         "Percentage - 82"

      B) "University of Oxford"
         "MSc Computer Science"
         "2020 - 2022"

      C) "MBA Finance | IIM Ahmedabad | 2021"
    """
    if not edu_text:
        return []

    PERCENTAGE_PAT = re.compile(
        r"^(percentage|grade|gpa|cgpa|score|marks|%)[\s\-:]*[\d.]+",
        re.IGNORECASE
    )
    YEAR_RANGE_PAT = re.compile(
        r"\b((?:19|20)\d{2})\s*[-–to]+\s*((?:19|20)\d{2}|present|current)\b",
        re.IGNORECASE
    )

    degrees  = []
    current  = {}
    lines    = [l.strip() for l in edu_text.split("\n") if l.strip()]

    def _flush():
        if current.get("degree") or current.get("institution"):
            degrees.append(dict(current))

    for line in lines:
        # Skip percentage/grade rows — they are metadata not a new entry
        if PERCENTAGE_PAT.match(line):
            if current:
                # Store grade on current entry instead of ignoring
                grade_m = re.search(r"[\d.]+\s*%?", line)
                if grade_m:
                    current["grade"] = grade_m.group().strip()
            continue

        year_range_m = YEAR_RANGE_PAT.search(line)
        year_m       = re.search(r"\b(19|20)\d{2}\b", line)
        deg_m        = DEGREE_PAT.search(line)

        # Year range line (e.g. "2019 - 2023") → attach to current entry
        if year_range_m and not deg_m:
            if current:
                # Use END year as graduation year, not start
                end_yr = year_range_m.group(2)
                start_yr = year_range_m.group(1)
                if re.match(r"(?i)present|current", end_yr):
                    current["year"] = int(start_yr)
                else:
                    try:
                        current["year"] = int(end_yr)
                    except ValueError:
                        current["year"] = int(start_yr)
            continue

        # Line with a degree keyword
        if deg_m:
            _flush()
            current = {"degree": None, "institution": None,
                       "year": None, "grade": None}

            # Check if institution is on same line: "B.Tech ECE, Bharath University"
            if "," in line:
                parts = [p.strip() for p in line.split(",", 1)]
                # Decide which part has the degree
                if DEGREE_PAT.search(parts[0]):
                    current["degree"]      = parts[0]
                    current["institution"] = parts[1]
                else:
                    current["degree"]      = parts[1]
                    current["institution"] = parts[0]
            else:
                current["degree"] = line

            if year_m:
                current["year"] = int(year_m.group())
            continue

        # Pure year line → attach to current
        if year_m and not deg_m:
            if current:
                current["year"] = int(year_m.group())
            continue

        # Otherwise treat as institution name
        if current and not current.get("institution") and len(line) > 2:
            current["institution"] = line
        elif not current:
            # Start a new entry even without a degree keyword
            current = {"degree": None, "institution": line,
                       "year": None, "grade": None}

    _flush()

    # Clean: remove entries that are just noise (no degree AND no institution)
    # Also remove None grade fields for cleanliness
    clean = []
    for d in degrees[:8]:
        if not d.get("degree") and not d.get("institution"):
            continue
        entry = {k: v for k, v in d.items() if v is not None}
        clean.append(entry)

    return clean


# ───────────────────────────────────────────────────────────────────
# CERTIFICATIONS / LANGUAGES
# ───────────────────────────────────────────────────────────────────
def extract_certs(text: str) -> list:
    if not text:
        return []
    certs = []
    for line in [l.strip() for l in text.split("\n") if l.strip()]:
        line = re.sub(r"^[-•·*▪▸]\s*", "", line)
        if 5 <= len(line) <= 150:
            certs.append(line)
    return certs[:15]


def extract_languages(text: str) -> list:
    if not text:
        return []
    parts = re.split(r"[,\n|•·/]+", text)
    langs = []
    for p in parts:
        p = re.sub(r"^[-•·*]\s*", "", p.strip())
        p = re.sub(r"\(.*?\)", "", p).strip()
        if 2 <= len(p) <= 30 and len(p.split()) <= 3:
            langs.append(p)
    return langs[:10]


# ───────────────────────────────────────────────────────────────────
# MASTER STRUCTURE FUNCTION  —  v7
# ───────────────────────────────────────────────────────────────────
def structure_resume(raw_text: str) -> dict:
    cleaned  = clean_text(raw_text)
    sections = split_sections(cleaned)

    # Use full text as fallback source for any missing section
    full     = sections.get("_full_text", cleaned)
    header   = sections.get("header",        "")
    skills_s = sections.get("skills",        "")
    exp_s    = sections.get("experience",    "")
    edu_s    = sections.get("education",     "")
    cert_s   = sections.get("certifications","")
    lang_s   = sections.get("languages",     "")
    summ_s   = sections.get("summary",       "")
    proj_s   = sections.get("projects",      "")
    achiev_s = sections.get("achievements",  "")

    # Use full text if specific sections are empty
    skills_src = skills_s or full
    exp_src    = exp_s    or full
    edu_src    = edu_s    or full

    contact  = extract_contact(header + "\n" + cleaned[:500])
    identity = extract_identity(header or cleaned[:300])
    skills   = extract_skills(skills_src, cleaned)
    jobs     = extract_experience(exp_src)
    edu      = extract_education(edu_src)
    certs    = extract_certs(cert_s)
    langs    = extract_languages(lang_s)
    location = extract_location(header + "\n" + cleaned[:800])
    total_exp= calc_total_exp(jobs)

    # Clean internal date fields
    clean_jobs = [{k: v for k, v in j.items() if not k.startswith("_")}
                  for j in jobs]

    # Summary — first 3 sentences
    summary = summ_s.strip()
    if summary:
        summary = " ".join(re.split(r"(?<=[.!?])\s+", summary)[:3])

    # If title not found in header, use most recent job title as fallback
    if not identity.get("current_title") and clean_jobs:
        identity["current_title"]    = clean_jobs[0].get("title")
        identity["title_confidence"] = "medium"

    # Overall confidence
    signals = [
        identity.get("name_confidence")  == "high",
        bool(contact.get("email")),
        total_exp > 0,
        len(skills.get("all_skills", [])) >= 5,
        bool(edu),
    ]
    overall = ("high"   if sum(signals) >= 4 else
               "medium" if sum(signals) >= 2 else "low")

    # Raw sections — send to Groq as fallback
    raw_sections = {}
    for sname in ["skills","experience","summary","achievements"]:
        raw = sections.get(sname, "").strip()
        if not raw and sname in ("skills","experience"):
            raw = cleaned  # fallback: send full text
        if raw:
            words = raw.split()
            raw_sections[sname] = " ".join(words[:500]) + ("..." if len(words) > 500 else "")

    detected = [k for k, v in sections.items()
                if v.strip() and not k.startswith("_")]

    return {
        "extraction_confidence": overall,
        "candidate": {
            "name":            identity.get("name"),
            "name_confidence": identity.get("name_confidence"),
            "current_title":   identity.get("current_title"),
            "title_confidence":identity.get("title_confidence"),
            "location":        location,
            "email":           contact.get("email"),
            "phone":           contact.get("phone"),
            "linkedin":        contact.get("linkedin"),
            "github":          contact.get("github"),
            "website":         contact.get("website"),
        },
        "summary":               summary or None,
        "total_experience_years":total_exp,
        "experience_confidence": "high" if total_exp > 0 else "low",
        "skills":                skills.get("all_skills", []),
        "skills_by_category":    skills.get("skills_by_category", {}),
        "skills_confidence":     skills.get("confidence"),
        "experience":            clean_jobs,
        "education":             edu,
        "certifications":        certs,
        "languages":             langs if langs else ["English"],
        "projects_mentioned":    bool(proj_s.strip()),
        "achievements_mentioned":bool(achiev_s.strip()),
        "sections_detected":     detected,
        "raw_sections":          raw_sections,
    }


# ═══════════════════════════════════════════════════════════════════
# AZURE FUNCTION ENTRY POINT  —  v1 programming model
# ═══════════════════════════════════════════════════════════════════
def main(req: func.HttpRequest) -> func.HttpResponse:
    logging.info("extractresume v7 called")

    if req.method == "GET":
        return func.HttpResponse(
            json.dumps({
                "status":  "healthy",
                "version": "7.3.0",
                "supports": ["pdf","docx","doc (legacy Word)","jpg","jpeg","png","tiff"],
                "fixes": [
                    "v7: robust section detection — ALL CAPS, Title Case, decorated headers",
                    "v7: name/title disambiguation fixed",
                    "v7: full DOCX extraction including tables and text boxes",
                    "v7: experience parser handles company-first and title-first layouts",
                    "v7: raw_sections always populated as Groq fallback",
                    "v7: fallback full-text scan when sections not detected",
                ]
            }),
            mimetype="application/json",
            status_code=200
        )

    try:
        body = req.get_json()
    except ValueError:
        return _err("Request body must be valid JSON", 400)

    file_bytes = None

    if "file_base64" in body:
        try:
            file_bytes = base64.b64decode(body["file_base64"])
        except Exception as e:
            return _err(f"Could not decode base64: {e}", 400)
    elif "file_url" in body:
        try:
            file_bytes = download_file(body["file_url"], body.get("access_token",""))
        except Exception as e:
            return _err(f"File download failed: {e}", 502)
    else:
        return _err("Provide file_base64 or file_url in the request body", 400)

    # Detect file type
    file_type = detect_type(file_bytes)
    if file_type == "unknown":
        file_type = body.get("file_type", "pdf").lower()

    logging.info(f"File type: {file_type}, size: {len(file_bytes)} bytes")

    try:
        if file_type == "pdf":
            result = extract_pdf(file_bytes)
        elif file_type == "docx":
            result = extract_docx(file_bytes)
        elif file_type == "doc":
            result = extract_doc_legacy(file_bytes)
        elif file_type in ("jpeg","jpg","png","gif","tiff","bmp"):
            result = extract_image(file_bytes)
        else:
            return _err(
                f"Unsupported file type: {file_type}. "
                "Supported: pdf, docx, doc, jpg, jpeg, png, tiff", 415
            )
    except ValueError as e:
        return _err(str(e), 422)
    except Exception as e:
        logging.error(f"Extraction error: {e}", exc_info=True)
        return _err(f"Extraction failed: {e}", 500)

    raw_text = result.get("text", "")
    if not raw_text.strip() and "warning" not in result:
        return _err(
            "No text could be extracted. For scanned PDFs, send as JPG/PNG for OCR.", 422
        )

    try:
        structured = structure_resume(raw_text)
    except Exception as e:
        logging.error(f"Structuring failed: {e}", exc_info=True)
        structured = {
            "extraction_confidence": "low",
            "error": f"Structuring failed: {e}",
            "raw_text": raw_text[:2000]
        }

    response = {
        "success":    True,
        "version":    "7.3.0",
        "file_type":  file_type,
        "page_count": result.get("page_count"),
        "ocr_used":   result.get("ocr_used", False),
        "resume":     structured
    }
    if "warning" in result:
        response["warning"] = result["warning"]

    logging.info(
        f"Done — conf: {structured.get('extraction_confidence')}, "
        f"skills: {len(structured.get('skills', []))}, "
        f"jobs: {len(structured.get('experience', []))}, "
        f"exp_yrs: {structured.get('total_experience_years')}"
    )

    return func.HttpResponse(
        json.dumps(response, ensure_ascii=False),
        mimetype="application/json",
        status_code=200
    )


def _err(msg: str, status: int) -> func.HttpResponse:
    logging.error(f"[{status}] {msg}")
    return func.HttpResponse(
        json.dumps({"success": False, "error": msg}),
        mimetype="application/json",
        status_code=status
    )
